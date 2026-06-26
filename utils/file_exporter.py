"""
文件导出工具模块

支持将内容导出为以下格式：
1. Markdown (.md)
2. PDF (.pdf) - 需要 markdowntopdf 或 weasyprint 支持
3. Word (.docx) - 需要 python-docx 支持

在模拟模式下，所有导出功能均可正常工作（PDF/Word 导出会给出提示）。
"""

import os
import json
import logging
from pathlib import Path
from typing import Optional, Union
from datetime import datetime

logger = logging.getLogger(__name__)


class FileExporter:
    """
    文件导出器
    
    支持导出格式：
    - Markdown（原生支持，无额外依赖）
    - PDF（需要 markdowntopdf 或 weasyprint）
    - Word（需要 python-docx）
    """

    def __init__(self, output_dir: Optional[Union[str, Path]] = None):
        """
        初始化导出器
        
        参数:
            output_dir: 输出目录，默认为项目 exports/ 目录
        """
        if output_dir is None:
            # 默认使用项目根目录下的 exports 文件夹
            self.output_dir = Path(__file__).resolve().parent.parent / "exports"
        else:
            self.output_dir = Path(output_dir)
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"文件导出目录: {self.output_dir}")

    def export_md(self, content: str, filename: str, metadata: Optional[dict] = None) -> Path:
        """
        导出为 Markdown 文件
        
        参数:
            content: Markdown 格式的内容
            filename: 文件名（不含扩展名）
            metadata: 可选的元数据（会以 YAML Front Matter 形式写入）
        
        返回:
            导出文件的完整路径
        """
        filepath = self.output_dir / f"{filename}.md"
        
        # 构建文件内容
        file_content = ""
        if metadata:
            file_content += "---\n"
            for key, value in metadata.items():
                file_content += f"{key}: {value}\n"
            file_content += "---\n\n"
        
        file_content += content
        
        # 写入文件
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(file_content)
        
        logger.info(f"已导出 Markdown 文件: {filepath}")
        return filepath

    def export_json(self, data: Union[dict, list], filename: str, indent: int = 2) -> Path:
        """
        导出为 JSON 文件
        
        参数:
            data: 要导出的数据（字典或列表）
            filename: 文件名（不含扩展名）
            indent: JSON 缩进空格数
        
        返回:
            导出文件的完整路径
        """
        filepath = self.output_dir / f"{filename}.json"
        
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=indent)
        
        logger.info(f"已导出 JSON 文件: {filepath}")
        return filepath

    def export_pdf(
        self,
        content: str,
        filename: str,
        mock_mode: bool = True,
        theme: str = "default"
    ) -> Path:
        """
        导出为 PDF 文件
        
        参数:
            content: Markdown 格式的内容
            filename: 文件名（不含扩展名）
            mock_mode: 模拟模式（True 时生成占位文件）
            theme: PDF 主题样式
        
        返回:
            导出文件的完整路径
        
        注意:
            非模拟模式下需要安装 markdowntopdf 或 weasyprint
        """
        filepath = self.output_dir / f"{filename}.pdf"
        
        if mock_mode:
            # 模拟模式：生成一个提示文件
            logger.warning("模拟模式：PDF 导出为占位文件")
            placeholder = f"""# PDF 导出占位文件

文件名: {filename}.pdf
生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
主题: {theme}

实际内容请查看对应的 .md 文件。

---
以下内容为原始 Markdown：

{content[:500]}...
"""
            with open(filepath.with_suffix(".pdf.placeholder.txt"), "w", encoding="utf-8") as f:
                f.write(placeholder)
            logger.info(f"[模拟] PDF 占位文件已生成: {filepath.with_suffix('.pdf.placeholder.txt')}")
            return filepath

        # 真实模式：尝试使用可用库生成 PDF
        return self._export_pdf_real(content, filepath, theme)

    def _export_pdf_real(self, content: str, filepath: Path, theme: str) -> Path:
        """真实模式 PDF 导出（尝试多种后端）"""
        # 方式1: 使用 weasyprint
        try:
            from weasyprint import HTML
            from .text_utils import MarkdownHelper
            
            html_content = MarkdownHelper.md_to_html(content, use_external=True)
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Noto Sans CJK SC", "Microsoft YaHei", sans-serif; margin: 2cm; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                    h2 {{ color: #34495e; margin-top: 1.5em; }}
                    code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                    pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                    blockquote {{ border-left: 4px solid #3498db; padding-left: 15px; color: #555; }}
                </style>
            </head>
            <body>{html_content}</body>
            </html>
            """
            HTML(string=full_html).write_pdf(str(filepath))
            logger.info(f"已导出 PDF 文件（weasyprint）: {filepath}")
            return filepath
        except ImportError:
            logger.warning("weasyprint 未安装，尝试其他方式...")

        # 方式2: 使用 reportlab（简单文本）
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import A4
            
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            styles = getSampleStyleSheet()
            story = [Paragraph(line, styles["Normal"]) for line in content.split("\n") if line.strip()]
            doc.build(story)
            logger.info(f"已导出 PDF 文件（reportlab）: {filepath}")
            return filepath
        except ImportError:
            logger.warning("reportlab 未安装")

        logger.error("PDF 导出失败：请安装 weasyprint 或 reportlab")
        raise RuntimeError("PDF 导出需要安装 weasyprint 或 reportlab，请运行: pip install weasyprint")

    def export_word(self, content: str, filename: str, mock_mode: bool = True) -> Path:
        """
        导出为 Word (.docx) 文件
        
        参数:
            content: Markdown 格式的内容
            filename: 文件名（不含扩展名）
            mock_mode: 模拟模式（True 时生成占位文件）
        
        返回:
            导出文件的完整路径
        
        注意:
            非模拟模式下需要安装 python-docx
        """
        filepath = self.output_dir / f"{filename}.docx"
        
        if mock_mode:
            # 模拟模式：生成占位文件
            logger.warning("模拟模式：Word 导出为占位文件")
            placeholder = f"""Word 导出占位文件

文件名: {filename}.docx
生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

实际内容请查看对应的 .md 文件。
原始 Markdown 内容已保存在同目录下的 .md 文件中。
"""
            with open(filepath.with_suffix(".docx.placeholder.txt"), "w", encoding="utf-8") as f:
                f.write(placeholder)
            logger.info(f"[模拟] Word 占位文件已生成: {filepath.with_suffix('.docx.placeholder.txt')}")
            return filepath

        # 真实模式：使用 python-docx
        return self._export_word_real(content, filepath)

    def _export_word_real(self, content: str, filepath: Path) -> Path:
        """真实模式 Word 导出"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re
            
            doc = Document()
            
            # 解析 Markdown 并添加到 Word 文档
            lines = content.split("\n")
            for line in lines:
                # 标题
                heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
                if heading_match:
                    level = min(len(heading_match.group(1)), 4)  # Word 最多支持 Heading 1-4
                    title = heading_match.group(2).strip()
                    doc.add_heading(title, level=level)
                    continue
                
                # 普通段落
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            doc.save(str(filepath))
            logger.info(f"已导出 Word 文件（python-docx）: {filepath}")
            return filepath
            
        except ImportError:
            logger.error("Word 导出失败：请安装 python-docx，运行 pip install python-docx")
            raise RuntimeError("Word 导出需要安装 python-docx")

    def export(
        self,
        content: str,
        filename: str,
        formats: list[str] = None,
        mock_mode: bool = True,
        **kwargs
    ) -> dict:
        """
        批量导出为多种格式
        
        参数:
            content: 要导出的内容（Markdown 格式）
            filename: 文件名（不含扩展名）
            formats: 导出格式列表，可选 ["md", "pdf", "docx", "json"]
            mock_mode: 是否使用模拟模式
            **kwargs: 传递给具体导出方法的参数
        
        返回:
            {格式: 文件路径} 的字典
        """
        if formats is None:
            formats = ["md"]
        
        results = {}
        
        for fmt in formats:
            fmt = fmt.lower().strip()
            try:
                if fmt == "md":
                    results[fmt] = self.export_md(
                        content, filename,
                        metadata=kwargs.get("metadata")
                    )
                elif fmt == "pdf":
                    results[fmt] = self.export_pdf(
                        content, filename,
                        mock_mode=mock_mode,
                        theme=kwargs.get("theme", "default")
                    )
                elif fmt == "docx":
                    results[fmt] = self.export_word(
                        content, filename,
                        mock_mode=mock_mode
                    )
                elif fmt == "json":
                    # JSON 导出需要 data 参数
                    data = kwargs.get("data", {"content": content})
                    results[fmt] = self.export_json(data, filename)
                else:
                    logger.warning(f"不支持的导出格式: {fmt}")
            except Exception as e:
                logger.error(f"导出 {fmt} 格式失败: {e}")
                results[fmt] = None
        
        return results

    @staticmethod
    def get_supported_formats() -> list[str]:
        """获取当前环境支持的导出格式"""
        supported = ["md", "json"]
        
        try:
            import weasyprint  # noqa
            supported.append("pdf")
        except ImportError:
            pass
        
        try:
            import docx  # noqa
            supported.append("docx")
        except ImportError:
            pass
        
        return supported
